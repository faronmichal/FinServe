import streamlit as st
from openai import OpenAI
from docx import Document
import io
import json

# mock data (instead of a real crm)
# in a real scenario, this data would be fetched via api
mock_crm_data = {
    "TechBuild Ltd.": {
        "tax_id": "1234567890",
        "industry": "Machinery manufacturing",
        "annual_revenue_usd": 1000000,
        "credit_history": "No delays, operating loan paid off in 2022",
        "internal_risk_score": "Low"
    },
    "GastroNova Inc.": {
        "tax_id": "0987654321",
        "industry": "Gastronomy",
        "annual_revenue_usd": 300000,
        "credit_history": "Two delays up to 14 days in 2023",
        "internal_risk_score": "Medium"
    }
}

# helper functions
def generate_credit_memo(api_key, client_data, sales_notes):
    # sends a request to openai to generate a credit memo
    client = OpenAI(api_key=api_key)
    
    prompt = f"""
    You are a credit analyst at FinServe. Your task is to prepare
    a professional credit memo based on hard data from the system
    and notes from a meeting with the client
    
    Hard data from CRM:
    {json.dumps(client_data, indent=2, ensure_ascii=False)}
    
    Sales representative notes from the meeting:
    "{sales_notes}"
    
    Required document format:
    - Client profile summary
    - Purpose of financing (inferred from notes)
    - Risk analysis (strengths and weaknesses)
    - Recommendation (approve / reject / requires additional documents)
    
    Write concisely, in a formal tone, in English
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return response.choices[0].message.content

def create_word_document(content, company_name):
    # generates a .docx file ready for download
    doc = Document()
    doc.add_heading(f'Credit Memo - {company_name}', 0)
    doc.add_paragraph(content)
    
    # saving to memory buffer
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# user interface (streamlit)
st.set_page_config(page_title="AI powered document generator", layout="centered")

st.title("AI powered document generator")
st.markdown("This tool automates the creation of credit memos based on CRM data and sales notes")

# sidebar for api key
with st.sidebar:
    st.header("Configuration")
    api_key_input = st.text_input("Enter OpenAI API key", type="password")
    st.info("The key is not saved. It is only used for this session.")

# main form
st.subheader("Select client (CRM data)")
selected_company = st.selectbox("Choose a company:", list(mock_crm_data.keys()))
client_data = mock_crm_data[selected_company]

with st.expander("Preview CRM data"):
    st.json(client_data)

st.subheader("Meeting notes (Sales Department)")
sales_notes = st.text_area(
    "Paste unstructured meeting notes:",
    height=150,
    placeholder="Example: The client needs 50k USD for new ovens. Recently late with taxes by 5 days. They promise it is a one-time thing"
)

# generate action
if st.button("Generate credit memo", type="primary"):
    if not api_key_input:
        st.error("Please provide the API key in the sidebar")
    elif not sales_notes:
        st.warning("Please enter meeting notes")
    else:
        with st.spinner("AI is analyzing data and generating the document..."):
            try:
                # ai call
                generated_memo = generate_credit_memo(api_key_input, client_data, sales_notes)
                
                # display result
                st.success("Document generated successfully!")
                st.markdown("### Document preview:")
                st.write(generated_memo)
                
                # option to download as word
                docx_file = create_word_document(generated_memo, selected_company)
                st.download_button(
                    label="Download as Word file (.docx)",
                    data=docx_file,
                    file_name=f"Credit_memo_{selected_company.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"An error occurred while connecting to the API: {e}")
                
                
